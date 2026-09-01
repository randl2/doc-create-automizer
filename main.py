import io
import os
import subprocess
import sys
import tempfile
import zipfile

from docx import Document
import pandas as pd
import streamlit as st

# ================= 1. НАЛАШТУВАННЯ СТОРІНКИ =================
st.set_page_config(
    page_title="Docx to PDF Generator",
    page_icon="📄",
    layout="centered"
)

st.title("📄 Генератор документів у PDF з архівацією в ZIP")[cite: 2]
st.caption(
    "Завантажте Word-шаблон і Excel-таблицю, щоб отримати готовий архів із PDF-файлами."[cite: 2]
)


# ================= 2. ФУНКЦІЇ ОБРОБКИ ТЕКСТУ В DOCX =================
def clean_and_replace_run_text(paragraphs, company, materials):
    """Склеює розірвані плейсхолдери та підставляє значення зі збереженням форматування."""
    for p in paragraphs:
        # 1. Склеюємо розірвані Вордом шматки плейсхолдерів[cite: 2]
        i = 0
        while i < len(p.runs) - 1:
            if any(part in p.runs[i].text for part in ["Company", "Material", "_name"]):[cite: 2]
                p.runs[i].text += p.runs[i + 1].text[cite: 2]
                p._p.remove(p.runs[i + 1]._r)[cite: 2]
                continue
            i += 1

        # 2. Заміна тексту[cite: 2]
        for run in p.runs:[cite: 2]
            if "Company_name" in run.text:[cite: 2]
                run.text = run.text.replace("Company_name", company)[cite: 2]
            if "Material_name" in run.text:[cite: 2]
                run.text = run.text.replace("Material_name", materials)[cite: 2]


def replace_placeholders(doc_path, company, materials, output_docx):
    """Замінює плейсхолдери в основному тексті та в усіх таблицях документа."""
    doc = Document(doc_path)[cite: 2]

    # Заміна в основному тексті[cite: 2]
    clean_and_replace_run_text(doc.paragraphs, company, materials)[cite: 2]

    # Заміна в таблицях[cite: 2]
    for table in doc.tables:[cite: 2]
        for row in table.rows:[cite: 2]
            for cell in row.cells:[cite: 2]
                clean_and_replace_run_text(cell.paragraphs, company, materials)[cite: 2]

    doc.save(output_docx)[cite: 2]


def convert_docx_to_pdf(docx_path, output_pdf_path):
    """Універсальна конвертація: через docx2pdf (на Windows) або LibreOffice (на Linux / Streamlit Cloud)."""
    if sys.platform == "win32":
        from docx2pdf import convert
        convert(docx_path, output_pdf_path)[cite: 2]
    else:
        output_dir = os.path.dirname(output_pdf_path) or "."
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                docx_path,
                "--outdir",
                output_dir,
            ],
            check=True,
        )[cite: 1]
        generated_pdf = os.path.splitext(docx_path)[0] + ".pdf"
        if generated_pdf != output_pdf_path and os.path.exists(generated_pdf):
            if os.path.exists(output_pdf_path):
                os.remove(output_pdf_path)
            os.rename(generated_pdf, output_pdf_path)


# ================= 3. ЗАВАНТАЖЕННЯ ФАЙЛІВ =================
st.subheader("1. Завантаження шаблону та таблиці")[cite: 2]
col1, col2 = st.columns(2)[cite: 2]

with col1:
    uploaded_template = st.file_uploader(
        "Шаблон Word (.docx)", type=["docx"], key="docx_uploader"
    )[cite: 2]

with col2:
    uploaded_excel = st.file_uploader(
        "Таблиця даних (.xlsx)", type=["xlsx", "xls", "csv"], key="excel_uploader"
    )[cite: 2]

# ================= 4. ОБРОБКА ДАНИХ ТА ГЕНЕРАЦІЯ =================
if uploaded_template and uploaded_excel:
    try:
        if uploaded_excel.name.lower().endswith(".csv"):
            df = pd.read_csv(uploaded_excel)
        else:
            df = pd.read_excel(uploaded_excel)[cite: 2]

        df.columns = df.columns.astype(str).str.strip()[cite: 2]

        # Автопошук колонок
        company_col = next(
            (c for c in df.columns if any(k in c.lower() for k in ["company", "назва", "організац", "компан"])),
            None,
        )[cite: 2]
        material_col = next(
            (c for c in df.columns if any(k in c.lower() for k in ["material", "матеріал", "продукц"])),
            None,
        )[cite: 2]

        if not company_col:
            st.error(f"❌ У таблиці не знайдено колонки з компанією. Знайдені колонки: {list(df.columns)}")[cite: 2]
        else:
            df_clean = df.dropna(subset=[company_col]).copy()[cite: 2]
            st.write(f"📊 Знайдено рядків для формування документів: **{len(df_clean)}**")[cite: 2]
            st.dataframe(df_clean.head(5), use_container_width=True)[cite: 2]

            # ================= 5. ЗАПУСК ГЕНЕРАЦІЇ =================
            st.markdown("---")[cite: 2]
            if st.button("🚀 Згенерувати PDF та створити ZIP-архів", type="primary"):[cite: 2]
                progress_bar = st.progress(0)[cite: 2]
                status_text = st.empty()[cite: 2]
                log_box = st.container()[cite: 2]

                with tempfile.TemporaryDirectory() as temp_dir:[cite: 2]
                    template_path = os.path.join(temp_dir, "Template.docx")[cite: 2]
                    with open(template_path, "wb") as f:[cite: 2]
                        f.write(uploaded_template.getbuffer())[cite: 2]

                    generated_pdf_paths = [][cite: 2]
                    total = len(df_clean)[cite: 2]

                    for index, (_, row) in enumerate(df_clean.iterrows()):[cite: 2]
                        company = str(row[company_col]).strip()[cite: 2]
                        materials = (
                            str(row[material_col]).strip()
                            if material_col and pd.notna(row[material_col])
                            else ""
                        )[cite: 2]

                        safe_company = "".join(c for c in company if c.isalnum() or c in (" ", "_", "-")).rstrip()[cite: 2]
                        temp_docx = os.path.join(temp_dir, f"temp_{safe_company}.docx")[cite: 2]
                        output_pdf = os.path.join(temp_dir, f"Пропозиція партнерства {safe_company}.pdf")[cite: 2]

                        status_text.text(f"Обробка [{index + 1}/{total}]: {company}...")[cite: 2]

                        # 1. Заміна плейсхолдерів у DOCX[cite: 2]
                        replace_placeholders(template_path, company, materials, temp_docx)[cite: 2]

                        # 2. Конвертація у PDF
                        try:
                            convert_docx_to_pdf(temp_docx, output_pdf)
                            if os.path.exists(output_pdf):[cite: 2]
                                generated_pdf_paths.append((output_pdf, os.path.basename(output_pdf)))[cite: 2]
                                log_box.success(f"✅ Згенеровано: {os.path.basename(output_pdf)}")[cite: 2]
                        except Exception as e:
                            log_box.error(f"❌ Помилка для {company}: {e}")[cite: 2]
                        finally:
                            if os.path.exists(temp_docx):[cite: 2]
                                os.remove(temp_docx)[cite: 2]

                        progress_bar.progress((index + 1) / total)[cite: 2]

                    # ================= 6. ПАКУВАННЯ В ZIP-АРХІВ =================
                    if generated_pdf_paths:[cite: 2]
                        status_text.text("Пакування файлів у ZIP-архів...")[cite: 2]
                        zip_buffer = io.BytesIO()[cite: 2]

                        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:[cite: 2]
                            for pdf_file_path, arcname in generated_pdf_paths:[cite: 2]
                                zip_file.write(pdf_file_path, arcname=arcname)[cite: 2]

                        zip_buffer.seek(0)[cite: 2]
                        status_text.text("Готово!")[cite: 2]
                        st.balloons()[cite: 2]
                        st.success(f"🎉 Успішно створено {len(generated_pdf_paths)} PDF-файлів!")[cite: 2]

                        st.download_button(
                            label="📥 Завантажити всі PDF в одному ZIP-архіві",
                            data=zip_buffer,
                            file_name="Пропозиції_партнерства_PDF.zip",
                            mime="application/zip",
                            type="primary",
                        )[cite: 2]
                    else:
                        st.warning("⚠️ Жодного PDF-файлу не вдалося згенерувати.")[cite: 2]

    except Exception as e:
        st.error(f"Помилка зчитування файлів: {e}")[cite: 2]
else:
    st.info("👆 Будь ласка, завантажте Word-шаблон та таблицю для початку.")[cite: 2]